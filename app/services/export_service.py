"""Service for exporting questions to various file formats."""

import html as html_module
import json
import logging
import re
from enum import Enum
from io import BytesIO
from pathlib import Path
from typing import List, Optional

import markdown
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, PageBreak, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.colors import black

from app.db.models import Question

logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    """Supported export formats."""

    TXT = "txt"
    PDF = "pdf"
    DOCX = "docx"
    JSON = "json"


class ExportService:
    """Service for exporting questions to files."""

    def export_to_txt(
        self, questions: List[Question], include_solutions: bool = False
    ) -> str:
        """
        Export questions to plain text format.

        Args:
            questions: List of questions to export
            include_solutions: Whether to include solutions

        Returns:
            Formatted text content
        """
        lines = []
        lines.append("=" * 80)
        lines.append("EXAM QUESTIONS")
        lines.append("=" * 80)
        lines.append("")

        for idx, question in enumerate(questions, 1):
            lines.append(f"Question {idx}")
            lines.append("-" * 80)
            lines.append(question.question_text)
            lines.append("")

            if include_solutions and question.solution:
                lines.append("Solution:")
                lines.append(question.solution)
                lines.append("")

            lines.append("")

        return "\n".join(lines)

    def export_to_json(
        self, questions: List[Question], include_solutions: bool = False
    ) -> str:
        """
        Export questions to JSON format.

        Args:
            questions: List of questions to export
            include_solutions: Whether to include solutions

        Returns:
            JSON string
        """
        data = {
            "questions": [
                {
                    "id": q.id,
                    "question_text": q.question_text,
                    "solution": q.solution if include_solutions else None,
                    "metadata": q.question_metadata or {},
                    "created_at": q.created_at.isoformat() if q.created_at else None,
                }
                for q in questions
            ],
            "total": len(questions),
        }
        return json.dumps(data, indent=2, ensure_ascii=False)

    def _markdown_to_reportlab_html(self, markdown_text: str) -> str:
        """
        Convert markdown to HTML that ReportLab can render.
        Matches the formatting used in the chat (LatexRenderer).
        
        Args:
            markdown_text: Markdown formatted text
            
        Returns:
            HTML string compatible with ReportLab
        """
        # Pre-process: Ensure horizontal rules (---) are on their own lines
        # This helps markdown library convert them properly
        lines = markdown_text.split('\n')
        processed_lines = []
        for i, line in enumerate(lines):
            # Check if line is a horizontal rule (3+ dashes, asterisks, or underscores)
            if re.match(r'^(\-{3,}|\*{3,}|_{3,})\s*$', line):
                # Ensure it's treated as a horizontal rule
                processed_lines.append('---')
            else:
                processed_lines.append(line)
        markdown_text = '\n'.join(processed_lines)
        
        # Convert markdown to HTML with more extensions for better formatting
        html = markdown.markdown(
            markdown_text, 
            extensions=['fenced_code', 'tables', 'nl2br', 'sane_lists']
        )
        
        # IMPORTANT: Process code blocks FIRST, before any other processing
        # This prevents LaTeX and other processing from interfering with code content
        
        # Extract and protect code blocks (both <pre><code> and inline <code>)
        code_blocks = []
        code_block_placeholders = []
        
        # Process code blocks (<pre><code>...</code></pre>) first
        def replace_code_block(match):
            code_content = match.group(1)
            # HTML escape the code content to prevent LaTeX/HTML interpretation
            escaped_code = html_module.escape(code_content)
            # Convert newlines to <br/> tags to preserve line breaks
            escaped_code = escaped_code.replace('\n', '<br/>')
            # Store placeholder and actual content
            placeholder = f'__CODE_BLOCK_{len(code_blocks)}__'
            code_blocks.append(f'<br/><font face="Courier" size="10">{escaped_code}</font><br/>')
            code_block_placeholders.append(placeholder)
            return placeholder
        
        html = re.sub(r'<pre><code>(.*?)</code></pre>', replace_code_block, html, flags=re.DOTALL)
        
        # Process inline code (<code>...</code>) - but only if not already inside <pre>
        def replace_inline_code(match):
            code_content = match.group(1)
            # HTML escape inline code
            escaped_code = html_module.escape(code_content)
            # Store placeholder and actual content
            placeholder = f'__INLINE_CODE_{len(code_blocks)}__'
            code_blocks.append(f'<font face="Courier" size="10">{escaped_code}</font>')
            code_block_placeholders.append(placeholder)
            return placeholder
        
        html = re.sub(r'<code>(.*?)</code>', replace_inline_code, html)
        
        # ReportLab's Paragraph supports limited HTML tags
        # Convert headings to styled text (ReportLab supports <b> and <font>)
        html = re.sub(r'<h1>(.*?)</h1>', r'<br/><b><font size="18">\1</font></b><br/>', html, flags=re.DOTALL)
        html = re.sub(r'<h2>(.*?)</h2>', r'<br/><b><font size="16">\1</font></b><br/>', html, flags=re.DOTALL)
        html = re.sub(r'<h3>(.*?)</h3>', r'<br/><b><font size="14">\1</font></b><br/>', html, flags=re.DOTALL)
        html = re.sub(r'<h4>(.*?)</h4>', r'<br/><b><font size="12">\1</font></b><br/>', html, flags=re.DOTALL)
        html = re.sub(r'<h5>(.*?)</h5>', r'<br/><b>\1</b><br/>', html, flags=re.DOTALL)
        html = re.sub(r'<h6>(.*?)</h6>', r'<br/><b>\1</b><br/>', html, flags=re.DOTALL)
        
        # Convert bold (ReportLab supports <b>)
        html = re.sub(r'<strong>(.*?)</strong>', r'<b>\1</b>', html, flags=re.DOTALL)
        
        # Convert italic (ReportLab supports <i>)
        html = re.sub(r'<em>(.*?)</em>', r'<i>\1</i>', html, flags=re.DOTALL)
        
        # Convert horizontal rules to a line of ASCII dashes (not Unicode)
        # Use simple dashes that ReportLab can render reliably
        # Add extra spacing before and after for visual separation
        horizontal_rule = '<br/><br/>' + ('-' * 70) + '<br/><br/>'
        html = re.sub(r'<hr\s*/?>', horizontal_rule, html)
        
        # Also handle any remaining --- patterns that weren't converted
        # This is a fallback in case markdown library didn't convert them
        html = re.sub(r'<br/>\s*---\s*<br/>', horizontal_rule, html)
        
        # Convert lists - ReportLab doesn't support <ul>/<ol> well, so convert to plain text with bullets
        # Add proper spacing for list items
        html = re.sub(r'<li>(.*?)</li>', r'<br/>• \1', html, flags=re.DOTALL)
        html = re.sub(r'<ul>|</ul>|<ol>|</ol>', '', html)
        
        # Remove paragraph tags but keep content with line breaks
        html = re.sub(r'<p>(.*?)</p>', r'\1<br/>', html, flags=re.DOTALL)
        
        # Remove any remaining unsupported HTML tags but keep their content
        # Keep only supported tags: b, i, br, font
        html = re.sub(r'<(?!\/?(?:b|i|br|font)[\s>])[^>]+>', '', html)
        
        # Clean up extra whitespace and line breaks
        # First, normalize newlines
        html = re.sub(r'\n+', '<br/>', html)
        # Then clean up excessive line breaks (more than 2 consecutive)
        html = re.sub(r'(<br/>){3,}', '<br/><br/>', html)
        # Clean up line breaks around horizontal rules
        html = re.sub(r'<br/>+(' + re.escape('-' * 70) + ')<br/>+', r'<br/><br/>\1<br/><br/>', html)
        
        # Escape ampersands that aren't part of entities
        html = re.sub(r'&(?!amp;|lt;|gt;|quot;|#)', '&amp;', html)
        
        # Restore code blocks AFTER all other processing is complete
        # This ensures code content is not affected by any regex replacements
        for placeholder, code_html in zip(code_block_placeholders, code_blocks):
            html = html.replace(placeholder, code_html)
        
        return html

    def export_to_pdf(
        self, questions: List[Question], include_solutions: bool = False
    ) -> BytesIO:
        """
        Export questions to PDF format with markdown rendering.

        Args:
            questions: List of questions to export
            include_solutions: Whether to include solutions

        Returns:
            BytesIO buffer containing PDF content
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        story = []
        styles = getSampleStyleSheet()

        # Custom styles for better formatting
        heading1_style = styles["Heading1"]
        heading1_style.fontSize = 18
        heading1_style.spaceAfter = 12
        
        heading2_style = styles["Heading2"]
        heading2_style.fontSize = 14
        heading2_style.spaceAfter = 8
        
        heading3_style = styles["Heading3"]
        heading3_style.fontSize = 12
        heading3_style.spaceAfter = 6

        # Title
        title = Paragraph("EXAM QUESTIONS", heading1_style)
        story.append(title)
        story.append(Spacer(1, 12))

        for idx, question in enumerate(questions, 1):
            # Check if this is a mock exam (full exam content)
            is_mock_exam = question.question_metadata and question.question_metadata.get("is_mock_exam", False)
            
            if is_mock_exam:
                # For mock exams, render the full exam content with markdown
                exam_html = self._markdown_to_reportlab_html(question.question_text)
                # Use a style that allows HTML formatting
                from reportlab.lib.styles import ParagraphStyle
                exam_style = ParagraphStyle(
                    'MockExam',
                    parent=styles['Normal'],
                    fontSize=11,
                    leading=14,
                    spaceAfter=12,
                    allowWidows=1,
                    allowOrphans=1,
                )
                try:
                    exam_para = Paragraph(exam_html, exam_style)
                    story.append(exam_para)
                except Exception as e:
                    # Fallback to plain text if HTML parsing fails
                    logger.warning(f"Failed to parse HTML for mock exam, using plain text: {e}")
                    exam_para = Paragraph(question.question_text.replace('\n', '<br/>'), exam_style)
                    story.append(exam_para)
                story.append(Spacer(1, 12))
            else:
                # Regular question format
                # Question number
                q_num = Paragraph(f"Question {idx}", heading2_style)
                story.append(q_num)
                story.append(Spacer(1, 6))

                # Question text with markdown support
                q_html = self._markdown_to_reportlab_html(question.question_text)
                q_text = Paragraph(q_html, styles["Normal"])
                story.append(q_text)
                story.append(Spacer(1, 12))

                # Solution if included
                if include_solutions and question.solution:
                    sol_title = Paragraph("Solution:", heading3_style)
                    story.append(sol_title)
                    story.append(Spacer(1, 6))
                    sol_html = self._markdown_to_reportlab_html(question.solution)
                    sol_text = Paragraph(sol_html, styles["Normal"])
                    story.append(sol_text)
                    story.append(Spacer(1, 12))
                
                # Page references if available
                if question.question_metadata and question.question_metadata.get("page_references"):
                    page_refs = question.question_metadata.get("page_references", [])
                    if page_refs:
                        ref_text = "References: " + ", ".join([
                            f"{ref.get('source_file', 'unknown')} (page {ref.get('page', '?')})"
                            for ref in page_refs[:5]  # Limit to first 5 references
                        ])
                        if len(page_refs) > 5:
                            ref_text += f", and {len(page_refs) - 5} more"
                        ref_para = Paragraph(f"<i>{ref_text}</i>", styles["Normal"])
                        story.append(ref_para)
                        story.append(Spacer(1, 6))

            story.append(Spacer(1, 12))
            
            # Add page break between questions if not last
            if idx < len(questions):
                story.append(PageBreak())

        doc.build(story)
        buffer.seek(0)
        return buffer

    def export_to_docx(
        self, questions: List[Question], include_solutions: bool = False
    ) -> BytesIO:
        """
        Export questions to DOCX format.

        Args:
            questions: List of questions to export
            include_solutions: Whether to include solutions

        Returns:
            BytesIO buffer containing DOCX content
        """
        doc = Document()
        doc.add_heading("EXAM QUESTIONS", 0)

        for idx, question in enumerate(questions, 1):
            doc.add_heading(f"Question {idx}", level=1)
            doc.add_paragraph(question.question_text)

            if include_solutions and question.solution:
                doc.add_heading("Solution:", level=2)
                doc.add_paragraph(question.solution)

            doc.add_paragraph("")  # Empty line between questions

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def export_questions(
        self,
        questions: List[Question],
        format: ExportFormat,
        include_solutions: bool = False,
    ) -> tuple[bytes, str, str]:
        """
        Export questions to specified format.

        Args:
            questions: List of questions to export
            format: Export format (txt, pdf, docx, json)
            include_solutions: Whether to include solutions

        Returns:
            Tuple of (content bytes, content type, file extension)

        Raises:
            ValueError: If format is not supported
        """
        if format == ExportFormat.TXT:
            content = self.export_to_txt(questions, include_solutions).encode("utf-8")
            return content, "text/plain", "txt"
        elif format == ExportFormat.JSON:
            content = self.export_to_json(questions, include_solutions).encode("utf-8")
            return content, "application/json", "json"
        elif format == ExportFormat.PDF:
            buffer = self.export_to_pdf(questions, include_solutions)
            return buffer.read(), "application/pdf", "pdf"
        elif format == ExportFormat.DOCX:
            buffer = self.export_to_docx(questions, include_solutions)
            return (
                buffer.read(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "docx",
            )
        else:
            raise ValueError(f"Unsupported export format: {format}")
